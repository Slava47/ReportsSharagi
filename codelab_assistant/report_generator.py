"""Модуль генерации отчетов в формате Word (.docx).

Создает готовый документ с титульным листом, блок-схемой,
листингом кода, результатами тестирования и выводами.
"""

import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from pygments import highlight
from pygments.formatters import NullFormatter
from pygments.lexers import CppLexer, DelphiLexer, PythonLexer

from .profiles import load_profile


def _get_lexer(language):
    """Возвращает лексер Pygments для языка.

    Args:
        language: Язык программирования.

    Returns:
        Экземпляр лексера Pygments.
    """
    lexers = {
        "pascal": DelphiLexer(),
        "python": PythonLexer(),
        "cpp": CppLexer(),
    }
    return lexers.get(language, PythonLexer())


def _add_title_page(doc, profile, student_info, analysis):
    """Добавляет титульный лист.

    Args:
        doc: Документ docx.
        profile: Профиль преподавателя.
        student_info: Информация о студенте.
        analysis: Результаты анализа кода.
    """
    font_name = profile.get("font_name", "Times New Roman")
    font_size = profile.get("font_size", 14)
    title_config = profile.get("title_page", {})

    # Верхняя часть — информация об учебном заведении
    for field in ["university", "faculty", "department"]:
        value = title_config.get(field, "")
        if value:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            run.font.name = font_name
            run.font.size = Pt(font_size)

    # Отступ
    for _ in range(4):
        doc.add_paragraph()

    # Тип работы
    work_type = title_config.get("work_type", "Лабораторная работа")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(work_type)
    run.font.name = font_name
    run.font.size = Pt(font_size + 4)
    run.bold = True

    # Дисциплина
    discipline = title_config.get("discipline", "Программирование")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'по дисциплине "{discipline}"')
    run.font.name = font_name
    run.font.size = Pt(font_size)

    # Тема/Вариант
    variant = student_info.get("variant", "")
    if variant:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Вариант {variant}")
        run.font.name = font_name
        run.font.size = Pt(font_size)

    # Отступ
    for _ in range(4):
        doc.add_paragraph()

    # Информация о студенте — правая сторона
    student_name = student_info.get("name", "Студент")
    group = student_info.get("group", "")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Выполнил: {student_name}")
    run.font.name = font_name
    run.font.size = Pt(font_size)

    if group:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"Группа: {group}")
        run.font.name = font_name
        run.font.size = Pt(font_size)

    # Отступ
    for _ in range(3):
        doc.add_paragraph()

    # Год
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(datetime.now().year))
    run.font.name = font_name
    run.font.size = Pt(font_size)

    # Разрыв страницы
    doc.add_page_break()


def _add_purpose_section(doc, profile, analysis):
    """Добавляет раздел с целью работы.

    Args:
        doc: Документ docx.
        profile: Профиль преподавателя.
        analysis: Результаты анализа кода.
    """
    font_name = profile.get("font_name", "Times New Roman")
    font_size = profile.get("font_size", 14)

    heading = doc.add_heading("Цель работы", level=1)
    for run in heading.runs:
        run.font.name = font_name

    p = doc.add_paragraph()
    run = p.add_run(analysis["purpose"])
    run.font.name = font_name
    run.font.size = Pt(font_size)

    # Алгоритмы
    if analysis.get("algorithms"):
        p = doc.add_paragraph()
        algorithms_text = ", ".join(analysis["algorithms"])
        run = p.add_run(
            f"Используемые алгоритмические концепции: {algorithms_text}."
        )
        run.font.name = font_name
        run.font.size = Pt(font_size)

    p = doc.add_paragraph()
    run = p.add_run(
        f"Язык программирования: {analysis['language_display']}."
    )
    run.font.name = font_name
    run.font.size = Pt(font_size)


def _add_flowchart_section(doc, profile, flowchart_path):
    """Добавляет раздел с блок-схемой.

    Args:
        doc: Документ docx.
        profile: Профиль преподавателя.
        flowchart_path: Путь к файлу изображения блок-схемы.
    """
    font_name = profile.get("font_name", "Times New Roman")
    font_size = profile.get("font_size", 14)

    heading = doc.add_heading("Блок-схема алгоритма", level=1)
    for run in heading.runs:
        run.font.name = font_name

    if flowchart_path and os.path.exists(flowchart_path):
        doc.add_picture(flowchart_path, width=Cm(15))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Рис. 1. Блок-схема алгоритма")
        run.font.name = font_name
        run.font.size = Pt(font_size - 2)
        run.italic = True
    else:
        p = doc.add_paragraph()
        run = p.add_run(
            "Блок-схема не была сгенерирована. "
            "Убедитесь, что Graphviz установлен в системе."
        )
        run.font.name = font_name
        run.font.size = Pt(font_size)


def _add_listing_section(doc, profile, analysis):
    """Добавляет раздел с листингом кода.

    Args:
        doc: Документ docx.
        profile: Профиль преподавателя.
        analysis: Результаты анализа кода.
    """
    font_name = profile.get("font_name", "Times New Roman")
    font_size = profile.get("font_size", 14)
    code_font = profile.get("code_font_name", "Courier New")
    code_size = profile.get("code_font_size", 10)

    heading = doc.add_heading("Листинг программы", level=1)
    for run in heading.runs:
        run.font.name = font_name

    p = doc.add_paragraph()
    run = p.add_run(f"Файл: {analysis['filename']}")
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.italic = True

    # Добавляем код с подсветкой синтаксиса (упрощённый вариант)
    lexer = _get_lexer(analysis["language"])
    tokens = lexer.get_tokens(analysis["code"])

    # Определяем цвета для разных типов токенов
    from pygments.token import (
        Comment,
        Keyword,
        Literal,
        Name,
        Number,
        Operator,
        String,
    )

    token_colors = {
        Keyword: RGBColor(0, 0, 255),       # Синий
        Comment: RGBColor(0, 128, 0),        # Зелёный
        String: RGBColor(163, 21, 21),       # Тёмно-красный
        Literal: RGBColor(163, 21, 21),      # Тёмно-красный
        Number: RGBColor(9, 134, 88),        # Тёмно-зелёный
        Name.Function: RGBColor(128, 0, 0),  # Тёмно-красный
        Operator: RGBColor(0, 0, 0),         # Чёрный
    }

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    for token_type, token_value in tokens:
        run = p.add_run(token_value)
        run.font.name = code_font
        run.font.size = Pt(code_size)

        # Подбираем цвет
        color = RGBColor(0, 0, 0)  # По умолчанию чёрный
        for parent_type, parent_color in token_colors.items():
            if token_type is parent_type or token_type in parent_type:
                color = parent_color
                break

        run.font.color.rgb = color

        # Ключевые слова — жирным
        if token_type is Keyword or token_type in Keyword:
            run.bold = True


def _add_test_results_section(doc, profile, test_results):
    """Добавляет раздел с результатами тестирования.

    Args:
        doc: Документ docx.
        profile: Профиль преподавателя.
        test_results: Результаты выполнения тестов.
    """
    font_name = profile.get("font_name", "Times New Roman")
    font_size = profile.get("font_size", 14)

    heading = doc.add_heading("Результаты тестирования", level=1)
    for run in heading.runs:
        run.font.name = font_name

    if not test_results or not test_results.get("compiled"):
        error_msg = "Программа не была выполнена."
        if test_results and test_results.get("compile_error"):
            error_msg += f" Ошибка компиляции: {test_results['compile_error']}"
        elif test_results is None:
            error_msg += " Тестовые данные не были предоставлены."
        p = doc.add_paragraph()
        run = p.add_run(error_msg)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        return

    results = test_results.get("results", [])
    if not results:
        p = doc.add_paragraph()
        run = p.add_run("Тестовые данные не были предоставлены.")
        run.font.name = font_name
        run.font.size = Pt(font_size)
        return

    # Таблица результатов
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Заголовки
    headers = ["№ теста", "Входные данные", "Выходные данные", "Статус"]
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header_text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = font_name
                run.font.size = Pt(font_size - 2)

    # Данные
    for result in results:
        row = table.add_row()

        row.cells[0].text = str(result.get("test_number", ""))
        row.cells[1].text = result.get("input", "").strip() or "(нет)"
        row.cells[2].text = result.get("stdout", "").strip() or "(нет вывода)"

        if result.get("error"):
            row.cells[3].text = f"Ошибка: {result['error']}"
        elif result.get("returncode", 0) != 0:
            row.cells[3].text = "Ошибка выполнения"
        else:
            row.cells[3].text = "Успешно"

        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size - 2)

    # Подпись таблицы
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Таблица 1. Результаты тестирования программы")
    run.font.name = font_name
    run.font.size = Pt(font_size - 2)
    run.italic = True


def _add_conclusion_section(doc, profile, analysis, test_results):
    """Добавляет раздел с выводами.

    Args:
        doc: Документ docx.
        profile: Профиль преподавателя.
        analysis: Результаты анализа кода.
        test_results: Результаты выполнения тестов.
    """
    font_name = profile.get("font_name", "Times New Roman")
    font_size = profile.get("font_size", 14)

    heading = doc.add_heading("Выводы", level=1)
    for run in heading.runs:
        run.font.name = font_name

    conclusion = (
        f"В ходе выполнения лабораторной работы была разработана программа "
        f"на языке {analysis['language_display']}. "
    )

    if analysis.get("algorithms"):
        algorithms_text = ", ".join(analysis["algorithms"])
        conclusion += (
            f"При реализации были использованы следующие алгоритмические "
            f"концепции: {algorithms_text}. "
        )

    if test_results and test_results.get("compiled"):
        total = len(test_results.get("results", []))
        success = sum(
            1 for r in test_results.get("results", [])
            if r.get("returncode", -1) == 0 and not r.get("error")
        )
        conclusion += (
            f"Программа была протестирована с {total} набором(ами) тестовых данных. "
            f"Успешно пройдено тестов: {success} из {total}. "
        )

    conclusion += "Цель работы достигнута."

    p = doc.add_paragraph()
    run = p.add_run(conclusion)
    run.font.name = font_name
    run.font.size = Pt(font_size)


def generate_report(
    analysis,
    test_results=None,
    flowchart_path=None,
    student_info=None,
    profile_name="default",
    output_path=None,
):
    """Генерирует отчет в формате Word.

    Args:
        analysis: Результаты анализа кода (из analyzer.analyze_code).
        test_results: Результаты тестирования (из executor.run_tests).
        flowchart_path: Путь к изображению блок-схемы.
        student_info: Словарь с данными студента (name, group, variant).
        profile_name: Имя профиля преподавателя.
        output_path: Путь для сохранения .docx файла.

    Returns:
        Путь к сгенерированному файлу.
    """
    if student_info is None:
        student_info = {"name": "Студент", "group": "", "variant": ""}

    profile = load_profile(profile_name)
    sections = profile.get("sections", [
        "title_page", "purpose", "flowchart", "listing",
        "test_results", "conclusion",
    ])

    doc = Document()

    # Настройка полей страницы
    margins = profile.get("margins", {})
    for section in doc.sections:
        section.top_margin = Cm(margins.get("top_cm", 2.0))
        section.bottom_margin = Cm(margins.get("bottom_cm", 2.0))
        section.left_margin = Cm(margins.get("left_cm", 3.0))
        section.right_margin = Cm(margins.get("right_cm", 1.5))

    # Генерируем разделы в указанном порядке
    if "title_page" in sections:
        _add_title_page(doc, profile, student_info, analysis)

    if "purpose" in sections:
        _add_purpose_section(doc, profile, analysis)

    if "flowchart" in sections:
        _add_flowchart_section(doc, profile, flowchart_path)

    if "listing" in sections:
        _add_listing_section(doc, profile, analysis)

    if "test_results" in sections:
        _add_test_results_section(doc, profile, test_results)

    if "conclusion" in sections:
        _add_conclusion_section(doc, profile, analysis, test_results)

    # Определяем путь сохранения
    if output_path is None:
        basename = os.path.splitext(analysis["filename"])[0]
        output_dir = os.path.join(os.path.dirname(__file__), "..", "reports")
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, f"report_{basename}.docx")

    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    return output_path
